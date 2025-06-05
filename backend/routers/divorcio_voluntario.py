#divorcio_voluntario.py
from fastapi import APIRouter, Form
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fastapi import Depends
from sqlalchemy.orm import Session
from database import get_db
from models import Documento
from routers.auth import obtener_usuario_actual  
import unicodedata
import uuid
import os
import locale
import datetime


router = APIRouter()

def normalizar(texto: str) -> str:
    if not texto:
        return ""
    return unicodedata.normalize("NFD", texto.strip().lower()).encode("ascii", "ignore").decode("utf-8")

@router.post("/generar/divorcio_voluntario")
async def generar_divorcio_voluntario(
    promovente1: str = Form(...),
    promovente2: str = Form(...),
    direccion_promovente: str = Form(...),
    fecha_matrimonio: str = Form(...),
    regimen_matrimonial: str = Form(...),
    cuantos_abogados: str = Form(...),
    abogados: str = Form(...),
    bienes_comunes: str = Form(None),
    total_bienes: str = Form(None),
    lista_bienes: str = Form(None),
    tiene_hijos: str = Form(...),
    hijos_info: str = Form(None),
    quien_guarda: str = Form(None),
    domicilio_hijos: str = Form(None),
    frecuencia_visitas: str = Form(None),
    horario_visitas: str = Form(None),
    porcentaje_alimentos: str = Form(None),
    uso_domicilio: str = Form(None),
    manutencion_conyuge: str = Form(None),
    conyuge_manutencion: str = Form(None),
    monto_manutencion: str = Form(None),
    usuario=Depends(obtener_usuario_actual),
    db: Session = Depends(get_db)
):
    locale.setlocale(locale.LC_TIME, 'es_MX.UTF-8')  # español de México
    ciudad = "Ciudad de México" 
    fecha_actual = datetime.datetime.now().strftime("%d de %B de %Y")
 

    doc = Document()
    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    encabezado_run = encabezado.add_run(
        f"{promovente1.upper()}\n"
        f"Vs\n"
        f"{promovente2.upper()}\n"
        f"JUICIO: DIVORCIO VOLUNTARIO CONTENCIOSO\n"
    )
    encabezado_run.font.size = Pt(16)

    doc.add_paragraph("\nC. JUEZ DE LO FAMILIAR EN TURNO DE PRIMERA INSTANCIA")
    doc.add_paragraph("DE LA CIUDAD DE MÉXICO")
    doc.add_paragraph("TRIBUNAL SUPERIOR DE JUSTICIA")

    abogado_lista = abogados.split(";")
    plural = len(abogado_lista) > 1
    if plural:
        texto_abogados = ", ".join([f"{a.split(':')[0]} (Cédula {a.split(':')[1]})" for a in abogado_lista])
        autorizacion = f"a los C.C. Licenciados en Derecho {texto_abogados}"
    else:
        nombre, cedula = abogado_lista[0].split(":")
        autorizacion = f"al C. Licenciado en Derecho {nombre} (Cédula {cedula})"

    doc.add_paragraph(
        f"P R E S E N T E:\n\n"
        f"{promovente1} y {promovente2}, por nuestro propio derecho, señalando como domicilio para oír y recibir toda clase de notificaciones, valores y documentos, "
        f"el ubicado en {direccion_promovente}, autorizando para tales efectos {autorizacion}, "
        f"ante Usted con el debido respeto comparecemos para exponer:\n\n"
        f"Que por medio del presente escrito, y con fundamento en los artículos 266, 267, 271, 272, 273, 282, 283 y 311 del Código Civil para la Ciudad de México, "
        f"y los artículos 1, 255, 256, 257 y demás relativos del Código de Procedimientos Civiles para la Ciudad de México, "
        f"venimos a promover JUICIO DE DIVORCIO VOLUNTARIO CONTENCIOSO, con base en los siguientes hechos y propuesta de convenio.\n"
    )
    # --- H E C H O S ---
    doc.add_heading("H E C H O S", level=1)

    doc.add_paragraph(
        f"1. Con fecha {fecha_matrimonio} los promoventes contrajimos matrimonio civil conforme a las leyes del Estado de la Ciudad de México, "
        f"lo que se acredita con el acta correspondiente que se exhibe.\n"
    )

    hijos_normalizado = normalizar(tiene_hijos)
    manutencion_normalizado = normalizar(manutencion_conyuge)
    hijos_lista = [h.strip() for h in hijos_info.split(";")] if hijos_info else []
    bienes_normalizado = normalizar(bienes_comunes)
    regimen_normalizado = normalizar(regimen_matrimonial)

    clausulas = []
    clausulas_num = 1
    numerales = ["PRIMERA", "SEGUNDA", "TERCERA", "CUARTA", "QUINTA", "SEXTA", "SÉPTIMA", "OCTAVA", "NOVENA", "DÉCIMA"]

    if hijos_normalizado == "si":
        if ";" in hijos_info:
            hijos_format = []
            for hijo in hijos_info.split(";"):
                nombre, edad = hijo.strip().split(":")
                hijos_format.append(f"{nombre} de {edad} años")
            hijos_texto = "; ".join(hijos_format)
            doc.add_paragraph(
                f"2. De dicho matrimonio procreamos a {hijos_texto}, quienes actualmente son menores de edad y se encuentran bajo nuestra responsabilidad y cuidado.\n")
        else:
            nombre_unico, edad = hijos_info.split(":")
            doc.add_paragraph(
                f"2. De dicho matrimonio procreamos a {nombre_unico}, quien actualmente cuenta con {edad} años de edad y se encuentra bajo nuestra responsabilidad y cuidado.\n")

        doc.add_paragraph(
            "3. Manifestamos nuestra voluntad de disolver el vínculo matrimonial mediante resolución judicial, ya que no se cumplen los requisitos del divorcio administrativo.\n")
        num_hecho = 4

    else:
        doc.add_paragraph("2. No procreamos hijos menores de edad ni existen personas incapaces a nuestro cargo, y ambas partes deseamos disolver el vínculo matrimonial de forma voluntaria ante la autoridad judicial.\n")
        num_hecho = 3

    doc.add_paragraph(
        f"{num_hecho}. Ambas partes presentamos junto a este escrito el convenio respectivo, mediante el cual se regulan las consecuencias personales y patrimoniales derivadas de la disolución del vínculo matrimonial.\n"
    )

    doc.add_heading("PROPUESTA DE CONVENIO", level=1)

    # --- Cláusulas por hijos ---
    if hijos_normalizado == "si":
        if len(hijos_lista) == 1:
            nombre_unico, edad = hijos_info.split(":")
            clausulas.append(
                f"{numerales[clausulas_num - 1]}.- La guarda y custodia de nuestro menor hijo {nombre_unico} de {edad} años quedará a cargo de {quien_guarda}, quien la ejercerá en el domicilio ubicado en {domicilio_hijos}.\n"
            )
        else:
            hijos_format = []
            for hijo in hijos_info.split(";"):
                nombre, edad = hijo.strip().split(":")
                hijos_format.append(f"{nombre} de {edad} años")
            hijos_texto = "; ".join(hijos_format)
            clausulas.append(
                f"{numerales[clausulas_num - 1]}.- La guarda y custodia de nuestros menores hijos {hijos_texto} quedará a cargo de {quien_guarda}, quien la ejercerá en el domicilio ubicado en {domicilio_hijos}.\n"
            )

        clausulas.append(
            f"{numerales[clausulas_num - 1]}.- El régimen de visitas y convivencias será ejercido por el progenitor que no tenga la custodia cada {frecuencia_visitas}, "
            f"en un horario de {horario_visitas}, procurando no afectar el desarrollo y bienestar de los menores.\n")
        clausulas_num += 1

        clausulas.append(
            f"{numerales[clausulas_num - 1]}.- En concepto de pensión alimenticia, el progenitor que no ejerza la custodia cubrirá el equivalente al {porcentaje_alimentos}% de sus ingresos, "
            f"destinado a cubrir alimentación, educación, salud, vestido y vivienda de los menores.\n")
        clausulas_num += 1

        clausulas.append(
            f"{numerales[clausulas_num - 1]}.- El uso del domicilio conyugal permanecerá a cargo de {uso_domicilio}, mientras los menores habiten con dicha persona, "
            f"con el objeto de preservar su estabilidad emocional y entorno habitual.\n")
        clausulas_num += 1

    # --- Cláusula de manutención conyugal ---
    if manutencion_normalizado == "si":
        clausulas.append(
            f"{numerales[clausulas_num - 1]}.- Se acuerda que {conyuge_manutencion} recibirá una pensión conyugal equivalente a {monto_manutencion}%, "
            f" conforme a lo dispuesto por la legislación vigente.\n")
        clausulas_num += 1

    # --- Cláusulas por bienes ---
    if bienes_normalizado == "si" and regimen_normalizado == "sociedad conyugal" and lista_bienes:
        bienes = lista_bienes.split(";")
        for bien in bienes:
            partes = bien.split(":")
            if len(partes) == 3:
                nombre, p1, p2 = partes
                clausulas.append(
                    f"{numerales[clausulas_num - 1]}.- En relación con el bien identificado como '{nombre}', se acuerda que "
                    f"{promovente1} recibirá el {p1}% y {promovente2} el {p2}%, quedando con ello concluida la distribución de dicho bien.\n"
                )
                clausulas_num += 1

    elif regimen_normalizado == "separacion de bienes":
        clausulas.append(
            f"{numerales[clausulas_num - 1]}.- Toda vez que el matrimonio se celebró bajo el régimen de separación de bienes, "
            f"cada parte conserva el dominio, uso y disfrute de los bienes que haya adquirido antes y durante el matrimonio.\n"
        )
        clausulas_num += 1


    # --- Si no hay nada que convenir ---
    if clausulas_num == 1:
        clausulas.append(
            "PRIMERA.- Manifestamos bajo protesta de decir verdad que no tenemos hijos menores, ni bienes que repartir, "
            "ni requerimos pensión alimenticia entre cónyuges, por lo que no resulta necesario convenir sobre estos aspectos.\n"
        )

    for cl in clausulas:
        doc.add_paragraph(cl)


    # --- D E R E C H O ---
    doc.add_heading("D E R E C H O", level=1)

    doc.add_paragraph(
        "En cuanto al fondo del asunto, son aplicables los artículos 266, 267, 271, 272, 273, 282, 283 y 311 del Código Civil para la Ciudad de México, "
        "así como los correlativos que regulan la disolución del vínculo matrimonial y sus efectos personales y patrimoniales.\n"
    )

    doc.add_paragraph(
        "El procedimiento se rige conforme a lo dispuesto por los artículos 1, 95, 255, 256, 257 y demás relativos del Código de Procedimientos Civiles "
        "para la Ciudad de México.\n"
    )

    # --- P R U E B A S ---
    doc.add_heading("P R U E B A S", level=1)

    doc.add_paragraph(
        "I.- LA CONFESIONAL.- A cargo de ambos promoventes, quienes deberán comparecer en forma personalísima a absolver posiciones al tenor del pliego correspondiente "
        "el día y hora que esta H. Autoridad señale, bajo apercibimiento de ley en caso de incomparecencia injustificada.\n"
    )

    doc.add_paragraph(
        "II.- LA DOCUMENTAL PÚBLICA.- Consistente en copia certificada del acta de matrimonio que se exhibe y acompaña al presente escrito.\n"
    )

    if hijos_normalizado == "si":
        doc.add_paragraph(
            "III.- LA DOCUMENTAL PÚBLICA.- Consistente en las actas de nacimiento de nuestros hijos menores, que se anexan en copia certificada para acreditar el vínculo filial y su edad.\n"
        )
    else:
        doc.add_paragraph(
            "III.- LA DOCUMENTAL PÚBLICA.- Consistente en comprobante de domicilio, que acredita la competencia territorial de este H. Juzgado para conocer del presente juicio.\n"
        )

    doc.add_paragraph(
        "IV.- LA INSTRUMENTAL DE ACTUACIONES.- Consistente en todas aquellas constancias procesales que obren en autos, así como las que se generen con motivo del presente procedimiento.\n"
    )

    doc.add_paragraph(
        "V.- LA PRESUNCIONAL LEGAL Y HUMANA.- En todo lo que favorezca a los intereses de los promoventes.\n"
    )

    doc.add_paragraph(
        "Todas las pruebas ofrecidas guardan relación directa con los hechos narrados y son conducentes para acreditar nuestras pretensiones.\n"
    )

    # --- P E T I C I O N E S ---
    doc.add_heading("P E T I C I O N E S", level=1)

    doc.add_paragraph(
        "Por lo anteriormente expuesto y fundado, a Usted C. Juez atentamente pedimos:\n\n"
        "PRIMERO.- Tenernos por presentados con este escrito, promoviendo JUICIO DE DIVORCIO VOLUNTARIO CONTENCIOSO.\n"
        "SEGUNDO.- Admitir la presente demanda y el convenio que se acompaña.\n"
        "TERCERO.- Señalar día y hora para la audiencia de ratificación del convenio.\n"
        "CUARTO.- Dictar sentencia definitiva que disuelva el vínculo matrimonial y apruebe el convenio en sus términos.\n"
        "QUINTO.- Ordenar al Registro Civil la anotación correspondiente en el acta de matrimonio.\n"
    )

    doc.add_paragraph(
        f"\nPROTESTAMOS LO NECESARIO.\n\n{ciudad}, a {fecha_actual}\n\n"
        f"___________________________________\n{promovente1.upper()}\n\n"
        f"_____________________________________\n{promovente2.upper()}"
    )

    excluir_justificacion = [
        "JUICIO: DIVORCIO VOLUNTARIO",
        "C. JUEZ DEL REGISTRO CIVIL DE LA CIUDAD DE MÉXICO.",
        "PROTESTAMOS LO NECESARIO."
    ]

        # Justificar todos los párrafos excepto los 3 primeros
    for p in doc.paragraphs:
        if not any(clave in p.text for clave in excluir_justificacion):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- GUARDAR DOCUMENTO Y RESPONDER ---
    nombre_archivo = f"Demanda_Divorcio_Voluntario_{promovente1.replace(' ', '_')}.docx"
    carpeta_usuario = f"documentos_usuario/{usuario.id}"
    os.makedirs(carpeta_usuario, exist_ok=True)
    ruta_completa = os.path.join(carpeta_usuario, nombre_archivo)
    doc.save(ruta_completa)

    nuevo_doc = Documento(
        usuario_id=usuario.id,
        nombre=nombre_archivo,
        ruta=ruta_completa
    )
    db.add(nuevo_doc)
    db.commit()

    return FileResponse(
        ruta_completa,
        filename=nombre_archivo,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
