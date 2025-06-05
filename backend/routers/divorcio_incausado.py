#divorcio_incausado.py
from fastapi import APIRouter, Form
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import unicodedata
import os
from fastapi import Depends
from sqlalchemy.orm import Session
from database import get_db
from models import Documento
from routers.auth import obtener_usuario_actual 
import locale
import datetime


router = APIRouter()

def normalizar(texto: str) -> str:
    if not texto:
        return ""
    return unicodedata.normalize("NFD", texto.strip().lower()).encode("ascii", "ignore").decode("utf-8")

@router.post("/generar/divorcio_incausado")
async def generar_divorcio_incausado(
    promovente: str = Form(...),
    demandado: str = Form(...),
    direccion_promovente: str = Form(...),
    cuantos_abogados: str = Form(...),
    abogados: str = Form(...),
    conoce_domicilio: str = Form(...),
    domicilio_demandado_si: str = Form(None),
    domicilio_demandado_no: str = Form(None),
    fecha_matrimonio: str = Form(...),
    regimen: str = Form(...),
    ultimo_domicilio: str = Form(...),
    fecha_separacion: str = Form(...),
    hijos: str = Form(...),
    hijos_info: str = Form(None),
    guarda_domicilio: str = Form(None),
    incluir_guardia: str = Form(None),
    guarda_titular: str = Form(None),
    incluir_alimentos: str = Form(None),
    visitas_frecuencia: str = Form(None),
    visitas_horario: str = Form(None),
    porcentaje_alimentos: str = Form(None),
    forma_pago_alimentos: str = Form(None),
    bienes: str = Form(None),
    lista_bienes: str = Form(None),
    proteccion: str = Form(None),
    usuario=Depends(obtener_usuario_actual),
    db: Session = Depends(get_db)
):
    locale.setlocale(locale.LC_TIME, 'es_MX.UTF-8')  # español de México
    fecha_actual = datetime.datetime.now().strftime("%d de %B de %Y")
    ciudad = "Ciudad de México"
    

    doc = Document()
    juicio_tipo = "DIVORCIO INCAUSADO"
    hijos_norm = normalizar(hijos)
    conoce_domicilio_norm = normalizar(conoce_domicilio)
    guardia_norm = normalizar(incluir_guardia)
    alimentos_norm = normalizar(incluir_alimentos)
    proteccion = proteccion or ""
    proteccion_norm = normalizar(proteccion)
    regimen_norm = normalizar(regimen)
    bienes_norm = normalizar(bienes)

    if hijos_norm == "si":
        if guardia_norm == "si" and alimentos_norm == "si":
            juicio_tipo += ", GUARDA, CUSTODIA Y ALIMENTOS"
        elif guardia_norm == "si":
            juicio_tipo += ", GUARDA Y CUSTODIA"
        elif alimentos_norm == "si":
            juicio_tipo += ", ALIMENTOS"


    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    encabezado_run = encabezado.add_run( f"{promovente.upper()}\nVs\n{demandado.upper()}\n JUICIO: {juicio_tipo}\n EXPEDIENTE: __________\n SECRETARÍA: __________")
    encabezado_run.font.size = Pt(16)

    doc.add_paragraph("\nC. JUEZ DE LO FAMILIAR EN TURNO DE PRIMERA INSTANCIA")
    doc.add_paragraph("DE LA CIUDAD DE MÉXICO")
    doc.add_paragraph("TRIBUNAL SUPERIOR DE JUSTICIA")

    abogado_lista = abogados.split(";")
    plural = len(abogado_lista) > 1
    if plural:
        texto_abogados = ", ".join([f"{a.split(':')[0]} (Cédula {a.split(':')[1]})" for a in abogado_lista])
        autorizacion = f"a los C.C. Licenciados en Derecho {texto_abogados}"
    else:
        nombre, cedula = abogado_lista[0].split(":")
        autorizacion = f"al C. Licenciado en Derecho {nombre} (Cédula {cedula})"

    doc.add_paragraph(
        f"P R E S E N T E:\n\n{promovente}, por propio derecho, señalando como domicilio para oír y recibir toda clase de notificaciones, valores y documentos, "
        f" el ubicado en {direccion_promovente}, {autorizacion}, ante Usted, con el debido respeto y consideración, comparezco y expongo:\n\n"
        f"Que por medio del presente escrito y con fundamento en los artículos 266 y 267 del Código Civil para la Ciudad de México, con reformas, "
        f"vengo a solicitar el divorcio por declaración unilateral.\n\n"
    )

    if conoce_domicilio_norm == "si":
        doc.add_paragraph(f"El promovente manifiesta conocer el domicilio del demandado, ubicado en {domicilio_demandado_si}.")
    else:
        doc.add_paragraph(
            f"Bajo protesta de decir verdad, desconozco su domicilio particular y para que pueda ser debidamente notificado señalo como domicilio ubicado en {domicilio_demandado_no}, "
            f"y toda vez que dicho domicilio se encuentra fuera de esta jurisdicción, solicito se gire atento exhorto al C. Juez competente de primera instancia para su emplazamiento y demás efectos legales."
        )

    # --- HECHOS ---
    doc.add_heading("H E C H O S", level=1)
    num_hecho = 1

    doc.add_paragraph(
        f"{num_hecho}. Con fecha {fecha_matrimonio}, el promovente contrajo matrimonio civil con el C. {demandado}, lo cual acredito con copia certificada del acta de matrimonio, "
        f"misma que se exhibe y se anexa para efectos legales correspondientes.\n"
    )
    num_hecho += 1

    regimen_legible = "Sociedad Conyugal" if regimen_norm == "sociedad conyugal" else "Separación de Bienes"
    doc.add_paragraph(
        f"{num_hecho}. Dicho vínculo matrimonial se celebró bajo el régimen de {regimen_legible}, como se acredita con el documento señalado.\n"
    )
    num_hecho += 1

    if hijos_norm == "si" and hijos_info:
        if ";" in hijos_info:
            hijos_format = []
            for hijo in hijos_info.split(";"):
                nombre, edad = hijo.strip().split(":")
                hijos_format.append(f"{nombre.strip()} de {edad.strip()} años")
            hijos_texto = "; ".join(hijos_format)
            doc.add_paragraph(
                f"{num_hecho}. Bajo protesta de decir verdad, de dicho matrimonio procreamos a {hijos_texto}, quienes actualmente son menores de edad, como se demuestra con las actas del Registro Civil que se anexan.\n"
            )
        else:
            nombre_unico, edad = hijos_info.strip().split(":")
            doc.add_paragraph(
                f"{num_hecho}. Bajo protesta de decir verdad, de dicho matrimonio procreamos a {nombre_unico.strip()} de {edad.strip()} años, quien actualmente es menor de edad, como se demuestra con el acta del Registro Civil que se anexa.\n"
            )
        num_hecho += 1

    doc.add_paragraph(
        f"{num_hecho}. Establecimos nuestro último domicilio conyugal en {ultimo_domicilio}, lo que da competencia a esta H. Autoridad.\n"
    )
    num_hecho += 1

    doc.add_paragraph(
        f"{num_hecho}. Desde {fecha_separacion}, nos encontramos separados de hecho, sin vida en común, situación que refleja la inexistencia de voluntad para continuar con el matrimonio.\n"
    )
    num_hecho += 1

    doc.add_paragraph(
        f"{num_hecho}. En virtud de lo anterior, y con fundamento en el artículo 266 del Código Civil para la Ciudad de México, solicito la disolución del vínculo matrimonial, en los términos del artículo 267 del mismo ordenamiento.\n"
    )
    num_hecho += 1

    doc.add_paragraph(
        f"{num_hecho}. A fin de dar cumplimiento a lo dispuesto por el artículo 267, presento la siguiente propuesta de convenio, con el objeto de regular las consecuencias inherentes a la disolución del vínculo matrimonial.\n"
    )
    # --- PROPUESTA DE CONVENIO ---
    doc.add_heading("PROPUESTA DE CONVENIO", level=1)
    clausula = 1

    if hijos_norm == "si" and guardia_norm == "si":
        if ";" in hijos_info:
            hijos_format = []
            for hijo in hijos_info.split(";"):
                nombre, edad = hijo.strip().split(":")
                hijos_format.append(f"{nombre.strip()} de {edad.strip()} años")
            hijos_texto = "; ".join(hijos_format)
            doc.add_paragraph(
                f"{clausula}.- La guarda y custodia de nuestros menores hijos {hijos_texto} quedará a cargo de {guarda_titular}, quien la ejercerá en el domicilio ubicado en {guarda_domicilio}, "
                f"procurando siempre el interés superior de los menores conforme a lo establecido en tratados internacionales y legislación nacional vigente.\n"
            )
            clausula += 1
            doc.add_paragraph(
                f"{clausula}.- El régimen de visitas y convivencias para el hoy demandado será cada {visitas_frecuencia}, recogiendo a los menores en el domicilio antes citado, en un horario de {visitas_horario}. "
                f"Dicho régimen se sujetará a condiciones que favorezcan el bienestar emocional y académico de los menores.\n"
            )
            clausula += 1
        else:
            nombre_unico, edad = hijos_info.strip().split(":")
            hijo_texto = f"{nombre_unico.strip()} de {edad.strip()} años"
            doc.add_paragraph(
                f"{clausula}.- La guarda y custodia de nuestro menor hijo {hijo_texto} quedará a cargo de {guarda_titular}, quien la ejercerá en el domicilio ubicado en {guarda_domicilio}, "
                f"procurando siempre el interés superior del menor conforme a lo establecido en tratados internacionales y legislación nacional vigente.\n"
            )
            clausula += 1
            doc.add_paragraph(
                f"{clausula}.- El régimen de visitas y convivencias para el hoy demandado será cada {visitas_frecuencia}, recogiendo al menor en el domicilio antes citado, en un horario de {visitas_horario}. "
                f"Dicho régimen se sujetará a condiciones que favorezcan el bienestar emocional y académico del menor.\n"
            )
            clausula += 1


    if hijos_norm == "si" and alimentos_norm == "si":
        doc.add_paragraph(
            f"{clausula}.- En cuanto a los alimentos, el hoy demandado deberá cubrir una pensión alimenticia del {porcentaje_alimentos}% de sus ingresos ordinarios y extraordinarios, "
            f"la cual deberá ser entregada de forma {forma_pago_alimentos}. Dicho porcentaje será destinado al sostenimiento integral de los menores conforme a lo previsto en los artículos 311 y siguientes del Código Civil para la Ciudad de México.\n"
        )
        clausula += 1

    if bienes_norm == "si" and lista_bienes:
        bienes = lista_bienes.split(";")
        for bien in bienes:
            nombre, pct1, pct2 = bien.split(":")
            doc.add_paragraph(
                f"{clausula}.- El bien identificado como '{nombre}' será distribuido en un {pct1}% a favor del promovente y un {pct2}% a favor del demandado.\n"
            )
            clausula += 1
    elif regimen_norm == "sociedad conyugal":
        doc.add_paragraph(f"{clausula}.- Bajo protesta de decir verdad, durante el matrimonio no se adquirieron bienes sujetos a reparto.\n")
        clausula += 1
    elif regimen_norm == "separacion de bienes":
        doc.add_paragraph(f"{clausula}.- Toda vez que el matrimonio se celebró bajo el régimen de separación de bienes, cada parte conserva el dominio, uso y disfrute de los bienes adquiridos antes y durante la vigencia del vínculo.\n ")
        clausula += 1
    # --- ORDEN DE PROTECCIÓN ---
    if proteccion_norm == "si":
        doc.add_heading("ORDEN DE PROTECCIÓN", level=1)
        doc.add_paragraph(
            f"Solicito se dicte orden de protección de carácter emergente en favor del promovente, consistente en requerir al hoy demandado para que se abstenga de acercarse "
            f"al domicilio particular, centro de trabajo o lugares públicos que frecuente el promovente, así como abstenerse de realizar cualquier acto de intimidación, hostigamiento o "
            f"agresión física, verbal o psicológica. De incumplirse lo anterior, solicito la intervención del Ministerio Público adscrito a este Juzgado, a efecto de que se integre la carpeta "
            f"de investigación correspondiente, conforme a lo establecido en la Ley de Acceso de las Mujeres a una Vida Libre de Violencia para la Ciudad de México.\n"
        )
    # --- DERECHO ---
    doc.add_heading("D E R E C H O", level=1)
    doc.add_paragraph(
        "En cuanto al FONDO del asunto, resultan aplicables los artículos 266, 267, 271, 282, 283 y 311, y demás relativos y aplicables del Código Civil para la Ciudad de México, con reformas vigentes.\n"
    )

    # --- PRUEBAS ---
    doc.add_heading("P R U E B A S", level=1)
    doc.add_paragraph(
        f"I.- LA CONFESIONAL.- A cargo del hoy demandado C. {demandado}, quien deberá absolver posiciones al tenor del pliego respectivo el día y hora que esta H. Autoridad señale, "
        f"debiendo comparecer de manera personalísima y no por conducto de apoderado o representante legal. Solicito que al momento del emplazamiento, también se le notifique sobre esta prueba, "
        f"bajo apercibimiento de ley en caso de incomparecencia injustificada.\n"
    )
    doc.add_paragraph(
        "II.- LA DOCUMENTAL PÚBLICA.- Consistente en la copia certificada del acta de matrimonio y, en su caso, las actas de nacimiento de los menores, las cuales se exhiben para acreditar "
        "los hechos manifestados en esta demanda.\n"
    )
    doc.add_paragraph(
        f"III.- LA DOCUMENTAL PÚBLICA.- Consistente en el comprobante de domicilio ubicado en {ultimo_domicilio}, como acreditación del último domicilio conyugal y competencia de este Juzgado.\n"
    )
    doc.add_paragraph(
        "IV.- LA INSTRUMENTAL DE ACTUACIONES.- Consistente en todas aquellas piezas procesales que obren en el expediente principal o que se integren con motivo del presente juicio.\n"
    )
    doc.add_paragraph(
        "V.- LA PRESUNCIONAL LEGAL Y HUMANA.- En todo lo que favorezca al promovente, con relación a los hechos descritos en el presente escrito.\n"
    )

    doc.add_paragraph(
        "Las pruebas que se ofrecen se relacionan directa y estrechamente con todos y cada uno de los hechos narrados y tienen por objeto acreditar los extremos de la acción ejercitada.\n"
    )
    # --- PETICIONES ---
    doc.add_heading("P E T I C I O N E S", level=1)
    doc.add_paragraph(
        "Por lo anteriormente expuesto y fundado, A Usted C. Juez atentamente pido se sirva:\n\n"
        "PRIMERO.- Tenerme por presentado con este escrito, promoviendo en la vía y forma legal el juicio de divorcio incausado.\n"
        "SEGUNDO.- Admitir la presente demanda y dar curso legal a la misma.\n"
        "TERCERO.- Girar exhorto o emitir orden de emplazamiento al hoy demandado en el domicilio indicado, para que comparezca y conteste lo que a su derecho convenga.\n"
        "CUARTO.- Tener por ofrecidas las pruebas señaladas, ordenando su admisión y desahogo en el momento procesal oportuno.\n"
        "QUINTO.- Dictar sentencia en la que se declare disuelto el vínculo matrimonial entre las partes.\n"
        "SEXTO.- Girar oficio al Registro Civil para que se haga la anotación correspondiente en el acta de matrimonio.\n"
    )

    excluir_justificacion = [
        "JUICIO: DIVORCIO ADMINISTRATIVO",
        "EXPEDIENTE: __________",
        "SECRETARÍA: __________",
        "C. JUEZ DEL REGISTRO CIVIL DE LA CIUDAD DE MÉXICO.",
        "P R E S E N T E:",
        "PROTESTAMOS LO NECESARIO."
    ]

        # Justificar todos los párrafos excepto los 3 primeros
    for p in doc.paragraphs:
        if not any(clave in p.text for clave in excluir_justificacion):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph(
        f"\nPROTESTO LO NECESARIO.\n\n{ciudad}, a {fecha_actual}\n\n _________________________________\n{promovente.upper()}"
    )

    nombre_archivo = f"Demanda_Divorcio_Incausado_{promovente.replace(' ', '_')}.docx"
    carpeta_usuario = f"documentos_usuario/{usuario.id}"
    os.makedirs(carpeta_usuario, exist_ok=True)
    ruta_completa = os.path.join(carpeta_usuario, nombre_archivo)
    doc.save(ruta_completa)

    nuevo_doc = Documento(
        usuario_id=usuario.id,
        nombre=nombre_archivo,
        ruta=ruta_completa
    )
    db.add(nuevo_doc)
    db.commit()

    return FileResponse(
        ruta_completa,
        filename=nombre_archivo,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
