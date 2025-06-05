from fastapi import APIRouter, Form, Depends
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from sqlalchemy.orm import Session
from database import get_db
from models import Documento
from routers.auth import obtener_usuario_actual
import os
import uuid
import datetime
import unicodedata

router = APIRouter()

def normalizar(texto: str) -> str:
    if not texto:
        return ""
    return unicodedata.normalize("NFD", texto.strip().lower()).encode("ascii", "ignore").decode("utf-8")

@router.post("/generar/guarda_custodia")
async def generar_guarda_custodia(
    promovente: str = Form(...),
    parentesco: str = Form(...),
    menores: str = Form(...),
    direccion_promovente: str = Form(...),
    cuantos_abogados: str = Form(...),
    abogados: str = Form(...),
    demandado: str = Form(...),
    conoce_domicilio: str = Form(...),
    domicilio_demandado: str = Form(None),
    domicilio_demandado_no: str = Form(None),
    tipo_relacion: str = Form(...),
    tiempo_convivencia: str = Form(...),
    motivo_guarda: str = Form(...),
    desea_visitas: str = Form(...),
    visitas: str = Form(None),
    restricciones: str = Form(None),
    usuario=Depends(obtener_usuario_actual),
    db: Session = Depends(get_db)
):
    ciudad = "Ciudad de Mexico"
    fecha = datetime.datetime.now().strftime("%d de %B de %Y")
    doc = Document()

    menores_lista = [m.strip() for m in menores.split(";") if m.strip()]
    plural = len(menores_lista) > 1
    menores_nombres = ", ".join([m.replace(":", " de ") + " anos" for m in menores_lista])
    menores_solo_nombres = ", ".join([m.split(":")[0] for m in menores_lista])

    abogado_lista = list({a.strip() for a in abogados.split(";") if a.strip()})
    plural_abogados = len(abogado_lista) > 1
    if plural_abogados:
        texto_abogados = ", ".join([f"{a.split(':')[0]} (Cedula {a.split(':')[1]})" for a in abogado_lista])
        autorizacion = f"a los C.C. Licenciados en Derecho {texto_abogados}"
    else:
        nombre, cedula = abogado_lista[0].split(":")
        autorizacion = f"al C. Licenciado en Derecho {nombre} (Cedula {cedula})"

    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    encabezado.add_run(
        f"{promovente.upper()}\nEn representacion de {menores_solo_nombres.upper()}\nVs\n{demandado.upper()}\n"
        f"JUICIO: GUARDA Y CUSTODIA\nEXPEDIENTE: __________\nSECRETARIA: __________"
    ).font.size = Pt(14)

    doc.add_paragraph("\nC. JUEZ DE LO FAMILIAR EN TURNO\nPODER JUDICIAL DE LA CIUDAD DE MEXICO\n")

    doc.add_paragraph(
        f"{promovente}, por propio derecho y en representacion de {'los menores' if plural else 'el menor'} {menores_nombres}, "
        f"senalando como domicilio para oir y recibir toda clase de notificaciones el ubicado en {direccion_promovente}, "
        f"{autorizacion}, ante Usted con el debido respeto comparezco y expongo:\n"
    )

    doc.add_paragraph(
        "Que por medio del presente escrito y con fundamento en los articulos 282, 283, 287, 288 y 291 del Codigo Civil para la Ciudad de Mexico, "
        "asi como el principio de interes superior del menor contenido en tratados internacionales, vengo a demandar la guarda y custodia "
        f"de {'los menores' if plural else 'el menor'} senalados.\n"
    )

    if normalizar(conoce_domicilio) == "si" and domicilio_demandado:
        doc.add_paragraph(f"El promovente manifiesta conocer el domicilio del demandado, ubicado en {domicilio_demandado}.")
    elif domicilio_demandado_no:
        doc.add_paragraph(
            f"Bajo protesta de decir verdad, manifiesto desconocer el domicilio actual del demandado, por lo que para efectos de notificacion "
            f"senalo como posible domicilio {domicilio_demandado_no}, y solicito se gire atento exhorto al C. Juez competente de primera instancia "
            f"en esa jurisdiccion para su legal emplazamiento y demas efectos legales a que haya lugar."
        )
    else:
        doc.add_paragraph("No se proporciono domicilio conocido ni alternativo para el demandado.")

    # Resto del documento (prestaciones, hechos, derecho, pruebas, peticiones) como antes
    # ...
        doc.add_heading("P R E S T A C I O N E S", level=1)
    doc.add_paragraph(f"1. Se me otorgue la guarda y custodia de {'los menores' if plural else 'el menor'}.")
    n = 2
    if normalizar(desea_visitas) == "si":
        doc.add_paragraph(f"{n}. Que se determine un régimen de convivencias adecuado entre {'los menores' if plural else 'el menor'} y el demandado.")
        n += 1
        if restricciones and normalizar(restricciones) == "si":
            doc.add_paragraph(f"{n}. Que se impongan restricciones a dichas convivencias por seguridad de {'los menores' if plural else 'el menor'}.")
            n += 1
    doc.add_paragraph(f"{n}. El pago de gastos y costas procesales.")

    doc.add_heading("H E C H O S", level=1)
    doc.add_paragraph(f"1. {menores_nombres} depende{'n' if plural else ''} económica y afectivamente de mí.")
    doc.add_paragraph(f"2. Soy {parentesco} de {'los menores' if plural else 'el menor'} y he asumido su cuidado, formación y protección.")
    doc.add_paragraph(f"3. Hubo convivencia entre las partes en calidad de {tipo_relacion} durante {tiempo_convivencia}.")
    doc.add_paragraph(f"4. Solicito la guarda y custodia por la siguiente razón: {motivo_guarda}.")
    if normalizar(desea_visitas) == "si" and visitas:
        doc.add_paragraph(f"5. Propongo el siguiente régimen de visitas: {visitas}.")
        if restricciones and normalizar(restricciones) == "si":
            doc.add_paragraph("6. Solicito que dichas convivencias se supervisen o limiten por antecedentes de riesgo.")
        doc.add_heading("D E R E C H O", level=1)
    doc.add_paragraph(
        "Con fundamento en los artículos 282, 283, 287, 288 y 291 del Código Civil para la Ciudad de México, "
        "y los artículos 255, 256, 260 y 261 del Código de Procedimientos Civiles para la CDMX, "
        "así como el principio de interés superior del menor consagrado en tratados internacionales."
    )

    doc.add_heading("P R U E B A S", level=1)
    doc.add_paragraph(f"1. Acta{'s' if plural else ''} de nacimiento de {'los menores' if plural else 'el menor'}.")
    doc.add_paragraph("2. Documentales que acreditan la relación y el domicilio.")
    doc.add_paragraph("3. Testigos sobre la convivencia y cuidados de los menores.")
    doc.add_paragraph("4. Presuncional legal y humana.")
    doc.add_paragraph("5. Instrumental de actuaciones.")

    doc.add_heading("P E T I C I O N E S", level=1)
    doc.add_paragraph("PRIMERO. Tenerme por presentado con esta demanda de guarda y custodia.")
    doc.add_paragraph("SEGUNDO. Admitirla y ordenar el emplazamiento del demandado.")
    doc.add_paragraph("TERCERO. Dictar sentencia otorgando la guarda y custodia al promovente.")
    if normalizar(desea_visitas) == "si":
        doc.add_paragraph("CUARTO. Fijar el régimen de convivencias propuesto, con o sin restricciones.")
    doc.add_paragraph("Último. Condenar al demandado al pago de costas procesales.")

    doc.add_paragraph(f"\nPROTESTO LO NECESARIO.\n{ciudad}, a {fecha}\n\n_______________________________\n{promovente.upper()}")
        

    excluir_justificacion = [
        "JUICIO: GUARDA Y CUSTODIA",
        "EXPEDIENTE: __________",
        "SECRETARIA: __________",
        "JUEZ DE LO FAMILIAR EN TURNO",
        "P R E S E N T E:",
        "PROTESTO LO NECESARIO."
    ]

    for p in doc.paragraphs:
        if not any(clave in p.text for clave in excluir_justificacion):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    nombre_archivo = f"Guarda_Custodia_{promovente.replace(' ', '_')}.docx"
    carpeta_usuario = f"documentos_usuario/{usuario.id}"
    os.makedirs(carpeta_usuario, exist_ok=True)
    ruta_completa = os.path.join(carpeta_usuario, nombre_archivo)
    doc.save(ruta_completa)

    nuevo_doc = Documento(
        usuario_id=usuario.id,
        nombre=nombre_archivo,
        ruta=ruta_completa
    )
    db.add(nuevo_doc)
    db.commit()

    return FileResponse(
        ruta_completa,
        filename=nombre_archivo,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
