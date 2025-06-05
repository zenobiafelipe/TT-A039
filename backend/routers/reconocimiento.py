from fastapi import APIRouter, Form
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fastapi import Depends
from sqlalchemy.orm import Session
from database import get_db
from models import Documento
from routers.auth import obtener_usuario_actual  
import datetime
import uuid
import os

router = APIRouter()

@router.post("/generar/reconocimiento_paternidad")
async def generar_reconocimiento_paternidad(
    promovente: str = Form(...),
    menor: str = Form(...),
    edad_menor: str = Form(...),
    fecha_nacimiento: str = Form(...),
    direccion: str = Form(...),
    demandado: str = Form(...),
    tipo_relacion: str = Form(...),
    periodo_relacion: str = Form(...),
    motivo: str = Form(...),
    conoce_trabajo: str = Form(...),
    trabajo: str = Form(None),
    direccion_trabajo: str = Form(None),
    ingreso: str = Form(None),
    domicilio: str = Form(...),
    solicita_pension: str = Form(...),
    porcentaje: str = Form(None),
    incumplimiento: str = Form(None),
    prueba_adn: str = Form(...),
    testigos: str = Form(None),
    usuario=Depends(obtener_usuario_actual),
    db: Session = Depends(get_db)
):
    fecha = datetime.datetime.now().strftime("%d de %B de %Y")
    ciudad = "Ciudad de México"
    doc = Document()

    header = doc.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header.add_run(
        f"{promovente.upper()}\nEn representación de {menor.upper()}\nVs\n{demandado.upper()}\n"
        f"JUICIO: RECONOCIMIENTO DE PATERNIDAD\nEXPEDIENTE: __________\nSECRETARÍA: __________"
    ).font.size = Pt(14)

    doc.add_paragraph("\nC. JUEZ DE LO FAMILIAR EN TURNO\nPODER JUDICIAL DE LA CIUDAD DE MÉXICO\nP R E S E N T E.\n")

    doc.add_paragraph(
        f"{promovente}, por mi propio derecho y en representación de la menor {menor}, quien actualmente tiene {edad_menor} años de edad, "
        f"señalando como domicilio para oír notificaciones {direccion}, comparezco y expongo:\n"
    )

    doc.add_heading("P R E S T A C I O N E S", level=1)
    doc.add_paragraph(f"1. El reconocimiento judicial de paternidad del C. {demandado} a favor de la menor {menor}.")
    if solicita_pension.lower() == "sí":
        doc.add_paragraph("2. El aseguramiento y fijación de una pensión alimenticia provisional y definitiva.")
    if prueba_adn.lower() == "sí":
        doc.add_paragraph("3. La práctica de prueba pericial en genética molecular (ADN).")
    doc.add_paragraph("4. El pago de costas procesales.")

    doc.add_heading("H E C H O S", level=1)
    doc.add_paragraph(f"1. La menor {menor} nació el {fecha_nacimiento}.")
    doc.add_paragraph(f"2. Sostuve una relación de tipo {tipo_relacion} con el demandado durante el periodo {periodo_relacion}.")
    doc.add_paragraph(f"3. Considero que es el padre de la menor porque {motivo}.")

    if conoce_trabajo.lower() == "sí":
        doc.add_paragraph(f"4. El demandado trabaja en {trabajo}, ubicado en {direccion_trabajo}, con un ingreso mensual aproximado de ${ingreso}.")
    else:
        doc.add_paragraph("4. Se desconoce el lugar de trabajo actual del demandado.")

    doc.add_paragraph(f"5. El demandado {domicilio} y no ha reconocido voluntariamente a la menor.")

    if solicita_pension.lower() == "sí":
        doc.add_paragraph(f"6. Solicito una pensión del {porcentaje} sobre sus ingresos, ya que ha incumplido desde {incumplimiento}.")

    doc.add_heading("D E R E C H O", level=1)
    doc.add_paragraph("Artículos 4º Constitucional, 361 al 380 y 391 del Código Civil para la CDMX, 255 y ss. del CPC local, Convención sobre los Derechos del Niño.")

    doc.add_heading("P R U E B A S", level=1)
    doc.add_paragraph("1. Acta de nacimiento de la menor.")
    if testigos:
        doc.add_paragraph(f"2. Testigos: {testigos}.")
    if prueba_adn.lower() == "sí":
        doc.add_paragraph("3. Prueba pericial en genética molecular (ADN).")
    doc.add_paragraph("4. Presuncional legal y humana. 5. Instrumental de actuaciones.")

    doc.add_heading("P E T I C I O N E S", level=1)
    doc.add_paragraph("PRIMERO. Tenerme por presentado con esta demanda y anexos.")
    doc.add_paragraph("SEGUNDO. Admitir el juicio y ordenar el emplazamiento del demandado.")
    doc.add_paragraph(f"TERCERO. Dictar sentencia que declare la paternidad del C. {demandado} respecto de la menor.")
    if solicita_pension.lower() == "sí":
        doc.add_paragraph("CUARTO. Fijar y asegurar la pensión alimenticia solicitada.")
    if prueba_adn.lower() == "sí":
        doc.add_paragraph("QUINTO. Girar oficio para la práctica de prueba pericial en ADN.")
    doc.add_paragraph("Último. Condenar al demandado al pago de costas del juicio.")

    doc.add_paragraph(f"\nPROTESTO LO NECESARIO.\n{ciudad}, a {fecha}\n\n{promovente.upper()}")

    
    excluir_justificacion = [
        "JUICIO: GUARDA Y CUSTODIA",
        "EXPEDIENTE: __________",
        "SECRETARIA: __________",
        "JUEZ DE LO FAMILIAR EN TURNO",
        "P R E S E N T E:",
        "PROTESTO LO NECESARIO."
    ]

    for p in doc.paragraphs:
        if not any(clave in p.text for clave in excluir_justificacion):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- Guardar y responder ---
    nombre_archivo = f"Reconocimiento_Paternidad_{promovente.replace(' ', '_')}.docx"
    carpeta_usuario = f"documentos_usuario/{usuario.id}"
    os.makedirs(carpeta_usuario, exist_ok=True)
    ruta_completa = os.path.join(carpeta_usuario, nombre_archivo)
    doc.save(ruta_completa)

    nuevo_doc = Documento(
        usuario_id=usuario.id,
        nombre=nombre_archivo,
        ruta=ruta_completa
    )
    db.add(nuevo_doc)
    db.commit()

    return FileResponse(
        ruta_completa,
        filename=nombre_archivo,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )