from fastapi import APIRouter, Form
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fastapi import Depends
from sqlalchemy.orm import Session
from database import get_db
from models import Documento
from routers.auth import obtener_usuario_actual  
import uuid
import datetime
import os

router = APIRouter()

@router.post("/generar/pension_alimenticia")
async def generar_pension_alimenticia(
    promovente: str = Form(...),
    parentesco: str = Form(...),
    direccion: str = Form(...),
    demandado: str = Form(...),
    menores: str = Form(...),
    ingresos: str = Form(...),
    cuantos_abogados: str = Form(...),
    abogados: str = Form(...),
    incumplimiento: str = Form(...),
    retroactivos: str = Form(...),
    medidas: str = Form(...),
    usuario=Depends(obtener_usuario_actual),
    db: Session = Depends(get_db)
):
    fecha = datetime.datetime.now().strftime("%d de %B de %Y")
    ciudad = "Ciudad de México"
    doc = Document()

    # --- Formateo menores ---
    lista_menores = [m.strip() for m in menores.split(";")]
    menores_format = []
    for menor in lista_menores:
        nombre, edad = [x.strip() for x in menor.split(":")]
        menores_format.append(f"{nombre} de {edad} años")
    menores_texto = ", ".join(menores_format[:-1]) + " y " + menores_format[-1] if len(menores_format) > 1 else menores_format[0]
    plural_menores = len(menores_format) > 1

    # --- Abogados ---
    abogado_lista = abogados.split(";")
    plural_abogados = len(abogado_lista) > 1
    if plural_abogados:
        texto_abogados = ", ".join([f"{a.split(':')[0]} (Cédula {a.split(':')[1]})" for a in abogado_lista])
        autorizacion = f"a los C.C. Licenciados en Derecho {texto_abogados}"
    else:
        nombre, cedula = abogado_lista[0].split(":")
        autorizacion = f"al C. Licenciado en Derecho {nombre} (Cédula {cedula})"

    # --- Encabezado ---
    header = doc.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header.add_run(
        f"{promovente.upper()}\nEn representación de {menores_texto.upper()}\nVs\n{demandado.upper()}\n"
        f"JUICIO: PENSIÓN ALIMENTICIA\nEXPEDIENTE: __________\nSECRETARÍA: __________"
    ).font.size = Pt(14)

    doc.add_paragraph("\nC. JUEZ DE LO FAMILIAR EN TURNO\nPODER JUDICIAL DE LA CIUDAD DE MÉXICO\nP R E S E N T E.\n")

    # --- Presentación con abogados ---
    doc.add_paragraph(
        f"{promovente}, por mi propio derecho y en representación de {'los menores' if plural_menores else 'el menor'} {menores_texto}, "
        f"señalando como domicilio para oír notificaciones {direccion}, {autorizacion}, comparezco para exponer:\n"
    )

    # --- Párrafo legal transitorio ---
    doc.add_paragraph(
        "Con fundamento en lo dispuesto por los artículos 4° de la Constitución Política de los Estados Unidos Mexicanos; 291, 303, "
        "308 al 323 Ter del Código Civil para la Ciudad de México; 255, 272 y demás relativos del Código de Procedimientos Civiles "
        "para la Ciudad de México; así como la Convención sobre los Derechos del Niño, vengo a promover formal demanda en la VÍA "
        "ORDINARIA CIVIL en la modalidad de PENSIÓN ALIMENTICIA."
    )

    # --- Prestaciones ---
    doc.add_heading("P R E S T A C I O N E S", level=1)
    doc.add_paragraph(f"1. El pago de una pensión alimenticia provisional y definitiva a favor de {'los menores' if plural_menores else 'el menor'}.")
    doc.add_paragraph(f"2. Fijación del monto conforme a las necesidades de {'los menores' if plural_menores else 'el menor'} y las posibilidades del demandado.")
    if retroactivos.lower() == "sí":
        doc.add_paragraph("3. El pago retroactivo desde la presentación de la demanda.")
    if medidas.lower() == "sí":
        doc.add_paragraph("4. Embargo precautorio u otra medida de aseguramiento.")
    doc.add_paragraph("5. El pago de costas procesales.")

    # --- Hechos ---
    doc.add_heading("H E C H O S", level=1)
    doc.add_paragraph(f"1. Soy {parentesco} de {'los menores' if plural_menores else 'el menor'} {menores_texto}, quienes dependen económicamente de mí.")
    doc.add_paragraph(f"2. El demandado {demandado} es el obligado alimentario, conforme a lo establecido en la ley.")
    doc.add_paragraph(f"3. {'Los menores requieren' if plural_menores else 'El menor requiere'} alimentos para cubrir sus necesidades: alimentación, educación, salud, vestido y vivienda.")
    doc.add_paragraph(f"4. El demandado tiene capacidad económica, con ingresos de aproximadamente ${ingresos} mensuales.")
    if incumplimiento.lower() == "sí":
        doc.add_paragraph("5. El demandado ha incumplido reiteradamente con su obligación alimentaria, a pesar de solicitudes previas.")
    else:
        doc.add_paragraph("5. El demandado no ha contribuido voluntariamente al sostenimiento de "
                          f"{'los menores' if plural_menores else 'el menor'}.")

    # --- Derecho ---
    doc.add_heading("D E R E C H O", level=1)
    doc.add_paragraph(
        "Fundamento en los artículos 4º de la Constitución, 291, 303, 308-323 Ter del Código Civil para la CDMX, "
        "y 255 y 272 del CPC para la CDMX, así como la Convención sobre los Derechos del Niño."
    )

    # --- Pruebas ---
    doc.add_heading("P R U E B A S", level=1)
    doc.add_paragraph(f"1. {'Las actas' if plural_menores else 'El acta'} de nacimiento de {'los menores' if plural_menores else 'el menor'}.")
    doc.add_paragraph("2. Comprobantes de gastos de manutención.")
    doc.add_paragraph("3. Prueba del ingreso del demandado.")
    doc.add_paragraph("4. Testigos, en su caso.")
    doc.add_paragraph("5. Documental pública, instrumental y presuncional legal y humana.")

    # --- Peticiones ---
    doc.add_heading("P E T I C I O N E S", level=1)
    doc.add_paragraph("PRIMERO. Tenerme por presentado con este escrito y copias, promoviendo demanda de pensión alimenticia.")
    doc.add_paragraph("SEGUNDO. Admitir la demanda en la vía correspondiente.")
    doc.add_paragraph("TERCERO. Dictar pensión provisional inmediata.")
    if retroactivos.lower() == "sí":
        doc.add_paragraph("CUARTO. Condenar al pago retroactivo desde esta fecha.")
    if medidas.lower() == "sí":
        doc.add_paragraph("QUINTO. Decretar embargo precautorio sobre los bienes del demandado.")
    doc.add_paragraph("Último. Dictar sentencia favorable y condenar al pago de costas.")

    # --- Firma ---
    doc.add_paragraph(f"\nPROTESTO LO NECESARIO.\n{ciudad}, a {fecha}\n\n_______________________________\n{promovente.upper()}")

    excluir_justificacion = [
        "JUICIO: PENSIÓN ALIMENTICIA",
        "EXPEDIENTE: __________",
        "SECRETARÍA: __________",
        "C. JUEZ DE LO FAMILIAR EN TURNO",
        "P R E S E N T E:",
        "PROTESTAMOS LO NECESARIO."
    ]

        # Justificar todos los párrafos excepto los 3 primeros
    for p in doc.paragraphs:
        if not any(clave in p.text for clave in excluir_justificacion):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- Guardar y responder ---
    nombre_archivo = f"Pension_Alimentaria_{promovente.replace(' ', '_')}.docx"
    carpeta_usuario = f"documentos_usuario/{usuario.id}"
    os.makedirs(carpeta_usuario, exist_ok=True)
    ruta_completa = os.path.join(carpeta_usuario, nombre_archivo)
    doc.save(ruta_completa)

    nuevo_doc = Documento(
        usuario_id=usuario.id,
        nombre=nombre_archivo,
        ruta=ruta_completa
    )
    db.add(nuevo_doc)
    db.commit()

    return FileResponse(
        ruta_completa,
        filename=nombre_archivo,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
