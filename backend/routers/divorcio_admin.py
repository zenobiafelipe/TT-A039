#divorcio_admin.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fastapi import APIRouter, Form
from fastapi.responses import FileResponse
from fastapi import Depends
from sqlalchemy.orm import Session
from database import get_db
from models import Documento
from routers.auth import obtener_usuario_actual 
import os
import locale
import datetime



router = APIRouter()

@router.post("/generar/divorcio_admin")
async def generar_divorcio_admin(
    promovente: str = Form(...),
    conyuge: str = Form(...),
    direccion: str = Form(...),
    fecha_matrimonio: str = Form(...),
    regimenadm: str = Form(...),
    usuario=Depends(obtener_usuario_actual),
    db: Session = Depends(get_db)

):
    locale.setlocale(locale.LC_TIME, 'es_MX.UTF-8')  # español de México
    fecha = datetime.datetime.now().strftime("%d de %B de %Y")
    ciudad = "Ciudad de México"
    doc = Document()

    # Encabezado
    header = doc.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = header.add_run(
        f"{promovente.upper()}\n"
        f"Vs\n"
        f"{conyuge.upper()}\n"
        f"JUICIO: DIVORCIO ADMINISTRATIVO\n"
    )
    run.font.size = Pt(16)

    doc.add_paragraph("\nC. JUEZ DEL REGISTRO CIVIL DE LA CIUDAD DE MÉXICO.\n")

    doc.add_paragraph(
        f"P R E S E N T E:\n\n"
        f"Quienes suscribimos, {promovente} y {conyuge}, por nuestro propio derecho, señalando como domicilio para oír y recibir notificaciones, valores y documentos, "
        f"el ubicado en {direccion}, comparecemos respetuosamente para exponer:\n\n"
        f"Que por medio del presente escrito, y con fundamento en el artículo 272 del Código Civil para la Ciudad de México, venimos a solicitar de manera conjunta y de común acuerdo "
        f"el divorcio por la vía administrativa, conforme a los siguientes:\n"
    )
    

    doc.add_heading("H E C H O S", level=1)
    doc.add_paragraph(f"1. Con fecha {fecha_matrimonio}, contrajimos matrimonio civil ante la autoridad correspondiente en la Ciudad de México, lo cual se acredita con el acta de matrimonio que anexamos al presente escrito.")
    doc.add_paragraph("2. A la fecha de presentación de esta solicitud, ambos comparecientes somos mayores de edad y manifestamos de forma libre, voluntaria y consciente nuestro deseo de disolver el vínculo matrimonial que nos une.")
    doc.add_paragraph("3. No hemos procreado hijos menores de edad, ni existen personas que dependan económicamente de nosotros.")
    doc.add_paragraph("4. La compareciente manifiesta bajo protesta de decir verdad, que no se encuentra en estado de gravidez.")
    doc.add_paragraph("5. Ninguno de los comparecientes requiere pensión alimenticia, lo que declaramos bajo protesta de decir verdad.")
    if regimenadm.lower() in ["sí", "si", "sociedad conyugal"]:
        doc.add_paragraph("6. Nuestro matrimonio se celebró bajo el régimen de sociedad conyugal, la cual fue previamente liquidada conforme a derecho.")
    elif regimenadm.lower() in ["no", "separación de bienes", "separacion de bienes"]:
        doc.add_paragraph("6. Nuestro matrimonio se celebró bajo el régimen de separación de bienes.")

    doc.add_heading("D E R E C H O", level=1)
    doc.add_paragraph( "En cuanto al fondo del presente asunto, resulta aplicable el artículo 272 del Código Civil para la Ciudad de México, así como las disposiciones relativas al divorcio administrativo establecidas en la legislación civil vigente." )
    doc.add_paragraph( "El procedimiento se tramita ante el C. Juez del Registro Civil, conforme a las formalidades establecidas en el mencionado numeral y demás correlativos aplicables.")

    doc.add_heading("P E T I C I O N E S", level=1)
    doc.add_paragraph(
        "Por lo anteriormente expuesto y fundado, a Usted C. Juez del Registro Civil atentamente solicitamos:\n\n"
        "PRIMERO.- Tenernos por presentados con este escrito en el que solicitamos la disolución del vínculo matrimonial que nos une, mediante el procedimiento de divorcio administrativo.\n"
        "SEGUNDO.- Que se tenga por acreditado que se reúnen todos y cada uno de los requisitos legales establecidos en el artículo 272 del Código Civil para la Ciudad de México.\n"
        "TERCERO.- Que se nos cite para que en el plazo legal acudamos a ratificar nuestra solicitud, conforme lo establece la ley.\n"
        "CUARTO.- Que una vez realizada la ratificación, se declare disuelto el vínculo matrimonial y se ordene hacer la anotación correspondiente en el acta de matrimonio.\n"
    )

    doc.add_paragraph(
        f"\nPROTESTAMOS LO NECESARIO.\n\n"
        f"{ciudad}, a {fecha}\n\n"
        f"_________________________________\n{promovente.upper()}\n\n"
        f"_________________________________\n{conyuge.upper()}"
    )
   
    excluir_justificacion = [
        "JUICIO: DIVORCIO ADMINISTRATIVO",
        "C. JUEZ DEL REGISTRO CIVIL DE LA CIUDAD DE MÉXICO.",
        "P R E S E N T E:",
        "PROTESTAMOS LO NECESARIO."
    ]

        # Justificar todos los párrafos excepto los 3 primeros
    for p in doc.paragraphs:
        if not any(clave in p.text for clave in excluir_justificacion):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Guardar el documento y devolverlo como archivo .docx
    nombre_archivo = f"Divorcio_Administrativo_{promovente.replace(' ', '_')}.docx"
    doc.save(nombre_archivo)


    # Opcional: guardar en carpeta por usuario
    carpeta_usuario = f"documentos_usuario/{usuario.id}"
    os.makedirs(carpeta_usuario, exist_ok=True)
    ruta_completa = os.path.join(carpeta_usuario, nombre_archivo)
    doc.save(ruta_completa)

    nuevo_doc = Documento(
        usuario_id=usuario.id,
        nombre=nombre_archivo,
        ruta=ruta_completa
    )
    db.add(nuevo_doc)
    db.commit()

    return FileResponse(ruta_completa, filename=nombre_archivo, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
